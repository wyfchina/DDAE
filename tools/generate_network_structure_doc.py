from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "网络结构评分V2数据结构与开发过程.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Pt(widths[index] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("网络结构评分 V2：数据结构与开发过程")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("开发原则：数据结构先行，算法其次，优化器最后。")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")
    p.add_run(f"\n{body}")


def add_table(doc: Document, headers, rows, widths) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_widths(table, widths)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.add_run("• ").bold = True
        p.add_run(item)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "核心边界",
        "解耦点在本系统中定义为物料节点。库存位置不决定解耦点，但决定该解耦点是否可执行：能否存放、由谁负责、是否可共享、质量状态和有效期如何控制。DDAE 负责评分、场景验证和主设置治理；PLM 管 BOM 结构；DDOM / SDBR 管详细执行。",
    )

    doc.add_heading("1. 总体路线", level=1)
    add_bullets(
        doc,
        [
            "数据结构先行：先把物料、BOM、供应、routing、库存位置和缓冲设置建成可解释网络。",
            "算法其次：先用图遍历和白盒评分识别候选控制点、库存缓冲、时间缓冲和能力缓冲。",
            "优化器最后：Gurobi / OR-Tools 只在候选动作很多且存在预算、能力、服务目标等组合约束时使用。",
            "评分建议不在同一次 DDMRP 推演中动态改变已生效解耦点；批准后的主设置才进入下一轮计划。",
        ],
    )

    doc.add_heading("2. 网络数据结构", level=1)
    add_table(
        doc,
        ["对象", "关键字段", "用途", "边界说明"],
        [
            ["ItemMaster", "itemCode, itemName, itemType, family, lifecycleStatus, unitCost", "定义成品、半成品、采购件和原材料节点。", "解耦点以 itemCode 为核心，不以工序或库存位置作为主定义。"],
            ["BomHeader", "bomId, parentItem, bomVersion, effectiveFrom, effectiveTo, releaseStatus", "承接 PLM 发布的 BOM 版本和生效期。", "DDAE 读取生效快照，不替代 PLM 管版本。"],
            ["BomLine", "parentItem, componentItem, quantityPer, scrapFactor, alternateGroup", "建立父子物料关系和用量传播。", "不同产品族用量差异由 BOM 行体现。"],
            ["AlternateItem", "alternateGroup, primaryItem, alternateItem, priority, substitutionRatio", "表达替代料路径和切换优先级。", "先做评分解释，后续再做替代供应场景。"],
            ["RoutingLine", "itemCode, modelCode, operationCode, resourceCode, capacityPerUnit, routingVersion", "把物料需求折算为资源负荷。", "工艺路线随型号变化，必须按型号/产品族和生效期管理。"],
            ["SupplierSource", "itemCode, supplierCode, allocationPercent, priority, leadTimeDays, leadTimeVariability, capacityPerWeek", "表达多供应商、主供/备供、交期和能力。", "lead time 固定，但用波动率表达不确定性。"],
            ["InventoryLocation", "itemCode, locationCode, locationType, qualityStatus, owner, shelfLifeDays", "判断缓冲点是否可执行。", "库存位置不是解耦点定义，但决定可存放性和责任归属。"],
            ["BufferSetting", "itemCode, isDecouplingPoint, inventoryBufferProfile, timeBufferDays, MOQ, orderCycleDays", "定义当前已生效的 DDMRP 主设置。", "当前 plan run 内固定，变更需进入主设置治理。"],
            ["LeadTimeProfile", "itemCode, standardLeadTimeDays, variabilityFactor, sourceType", "支持时间缓冲评分。", "可从采购、外协、前道工序或质量放行来源聚合。"],
            ["NetworkScoreResult", "candidateId, targetItem, score, settingType, rationale, evidence", "保存评分结果和解释。", "是评审入口，不是自动执行指令。"],
        ],
        [1300, 2500, 2800, 2760],
    )

    doc.add_heading("3. 网络图关系", level=1)
    add_table(
        doc,
        ["边类型", "方向", "用途"],
        [
            ["BOM 边", "parentItem -> componentItem", "展开物料结构、计算用量和下游需求传播。"],
            ["反向影响边", "componentItem -> parentItem", "计算某物料影响哪些成品、产品族和客户交付。"],
            ["供应边", "supplierCode -> itemCode", "识别多供应商、主供/备供和供应风险传播。"],
            ["Routing 边", "itemCode -> resourceCode", "把预计补货订单折算为资源负荷和能力缓冲候选。"],
            ["库存边", "itemCode -> locationCode", "判断控制点物料是否具备库存承载位置。"],
            ["替代料边", "itemCode <-> alternateItem", "评估替代路径、切换成本和供应风险缓解能力。"],
        ],
        [1800, 2600, 4960],
    )

    doc.add_heading("4. V2 白盒评分指标", level=1)
    add_table(
        doc,
        ["指标", "计算口径", "业务解释"],
        [
            ["下游覆盖度", "从物料节点反向遍历到成品 SKU / 产品族的数量。", "覆盖越广，越可能成为控制点或共用缓冲点。"],
            ["数量影响度", "下游需求 × BOM 用量 × 损耗因子聚合。", "不是只看出现次数，还要看真实需求拉动量。"],
            ["累计提前期", "物料节点到成品或供应节点路径上的 lead time 聚合。", "用于识别时间缓冲和解耦保护价值。"],
            ["供应风险", "供应商数量、主供风险、备供能力、lead time 波动率、供应能力缺口。", "决定是否优先加时间缓冲或供应主设置治理。"],
            ["资源约束影响", "下游补货需求 × routing capacityPerUnit 聚合到资源。", "决定是否需要能力缓冲或提前建库。"],
            ["库存代价", "单位成本、库存位置、质量状态、货架期、库存空间。", "成本高不一定不重要，但可能更适合时间缓冲。"],
            ["管理复杂度", "涉及产品族、责任组织、供应商、库存位置和替代路径数量。", "避免推荐难以执行的复杂控制点。"],
        ],
        [1600, 3900, 3860],
    )

    doc.add_heading("5. 候选类型与判断", level=1)
    add_table(
        doc,
        ["候选类型", "优先条件", "输出建议"],
        [
            ["解耦点 / 控制点", "下游覆盖广、累计提前期长、共用性高、能隔离供应或资源波动。", "建议将物料纳入 DDOM 主设置治理，评估是否成为缓冲物料。"],
            ["库存缓冲点", "需求稳定或可预测、库存代价可接受、具备可执行库存位置。", "建议设置或重审红黄绿区、MOQ、订货周期。"],
            ["时间缓冲点", "长交期、高波动率、高供应风险、高成本物料。", "建议在控制点物料前增加保护时间和 Act/Late 阈值。"],
            ["能力缓冲点", "routing 聚合后形成瓶颈、峰值负荷高、影响多个产品族。", "建议设置能力保护边界、增班/外协/窗口预留策略。"],
            ["只监控不建缓冲", "风险存在但库存代价高、替代路径明确或管理复杂度高。", "建议加入异常监控，不直接增加库存。"],
        ],
        [1800, 4200, 3360],
    )

    doc.add_heading("6. 详细开发过程", level=1)
    add_table(
        doc,
        ["阶段", "开发内容", "验收标准"],
        [
            ["Phase 1 数据模型", "新增 ItemMaster、BomHeader、BomLine、AlternateItem、RoutingLine、SupplierSource、InventoryLocation、BufferSetting、LeadTimeProfile DTO。", "模型能表达 PLM BOM 快照、多供应商、型号化 routing 和库存位置。"],
            ["Phase 2 Seed 数据", "构造卫星制造多层 BOM：卫星平台、有效载荷、星载电子、热控结构、FPGA、电缆束等。", "至少覆盖 4 个产品族、10 层以内示例路径、多供应商和替代料。"],
            ["Phase 3 图构建服务", "建立正向/反向邻接表，校验循环、孤儿节点、版本生效期和缺失主数据。", "能按物料展开上游/下游影响范围，并输出校验报告。"],
            ["Phase 4 网络指标计算", "计算下游覆盖度、数量影响度、累计提前期、供应风险、资源约束、库存代价。", "每个指标都能追溯到具体 BOM 行、供应来源或 routing 行。"],
            ["Phase 5 V2 评分服务", "把指标转成解耦点、库存缓冲、时间缓冲、能力缓冲候选评分。", "候选结果包含得分、推荐类型、证据、解释和不采纳风险。"],
            ["Phase 6 场景验证", "把候选作为拟议主设置，运行 DDS&OP / DDMRP 白盒推演。", "输出库存金额、红区周、补货订单、RCCP、供应缺口变化。"],
            ["Phase 7 主设置治理", "将通过评审的候选保存为主设置变更请求。", "批准后成为下一轮 plan run 输入，不在当前推演中动态改写。"],
            ["Phase 8 优化器组合选择", "当候选动作很多且存在预算/能力/服务目标约束时，调用 Gurobi / OR-Tools。", "优化器只选择动作组合，最终必须回到白盒引擎重算。"],
        ],
        [1400, 4800, 3160],
    )

    doc.add_heading("7. 计划员解释口径", level=1)
    add_bullets(
        doc,
        [
            "为什么推荐这个物料：说明它影响哪些成品、产品族、供应商、库存位置和关键资源。",
            "为什么推荐这种缓冲：说明更适合库存缓冲、时间缓冲、能力缓冲，还是只监控。",
            "如果采纳会怎样：展示库存金额、红区周、服务风险、资源负荷和供应缺口变化。",
            "如果不采纳会怎样：说明可能影响的成品、缺料周、瓶颈资源和供应商升级事项。",
            "评分解释应该面向业务，而不是只展示公式。公式可进入白盒追踪和审计链。",
        ],
    )

    doc.add_heading("8. 与现有系统的衔接", level=1)
    add_table(
        doc,
        ["系统", "职责"],
        [
            ["PLM", "管理 BOM 版本、生效日期、替代料和发布状态。"],
            ["DDAE / DDS&OP", "读取网络快照，做网络评分、场景预览、优化推荐和主设置治理。"],
            ["DDOM / SDBR", "执行详细资源日历、日能力、排程、工单释放和现场反馈。"],
            ["Gurobi / OR-Tools", "仅在候选动作组合选择层使用，不替代 DDMRP 规则和白盒推演。"],
        ],
        [1800, 7560],
    )

    doc.add_heading("9. 当前确认的业务原则", level=1)
    add_bullets(
        doc,
        [
            "BOM 版本、生效日期、替代料由 PLM 管理，DDAE 使用已发布快照。",
            "同一物料在不同产品族中的用量差异由 BOM 行体现。",
            "采购 lead time 使用固定值，波动率作为评分和时间缓冲因子。",
            "工艺路线会随型号变化，因此 routing 必须支持型号/产品族和版本。",
            "一个物料可以有多个供应商，评分需要区分主供、备供、能力和切换风险。",
            "解耦点定义为物料节点；库存位置是可执行性属性，不是解耦点定义本身。",
            "时间缓冲放在控制点物料前，用于保护控制点物料的供应、前道工序或质量放行风险。",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
